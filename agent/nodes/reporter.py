import os
import re
from agent.llm_factory import get_llm
from agent.state import AnalysisState
import config
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if p.runs:
        if level == 1:
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            p.runs[0].font.color.rgb = RGBColor(0x2C, 0x52, 0x82)


def _add_paragraph(doc, text, size=11):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    if p.runs:
        p.runs[0].font.size = Pt(size)


def _add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_table_from_text(doc, stdout):
    lines = [l for l in stdout.strip().split("\n") if l.strip()]
    if len(lines) < 2:
        doc.add_paragraph(stdout[:400], style="No Spacing")
        return

    def split_row(line):
        parts = re.split(r"\s{2,}", line.strip())
        return [p.strip() for p in parts if p.strip()]

    headers = split_row(lines[0])
    if not headers:
        doc.add_paragraph(stdout[:400], style="No Spacing")
        return

    rows = []
    for line in lines[1:30]:
        row = split_row(line)
        if row:
            rows.append(row)

    if not rows:
        doc.add_paragraph(stdout[:400], style="No Spacing")
        return

    max_cols = max(len(headers), max(len(r) for r in rows))
    while len(headers) < max_cols:
        headers.insert(0, "")
    for row in rows:
        while len(row) < max_cols:
            row.append("")
        row[:] = row[:max_cols]

    table = doc.add_table(rows=1 + len(rows), cols=max_cols)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers[:max_cols]):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        _shade_cell(hdr_cells[i], "D6E4F0")

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row[:max_cols]):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            if r_idx % 2 == 1:
                _shade_cell(cells[c_idx], "F7FAFC")

    doc.add_paragraph("")


def _embed_chart(doc, chart_path):
    full_path = os.path.join(config.CHARTS_FOLDER, chart_path)
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")


def _ask_llm(llm, prompt_text, max_chars=1200):
    """llm is passed in explicitly — no global state."""
    response = llm.invoke(prompt_text)
    text = response.content.strip()
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"#+\s*", "", text)
    return text[:max_chars]


def _format_findings_text(findings):
    parts = []
    for i, f in enumerate(findings, 1):
        parts.append(
            f"Finding {i} [{f.get('status','?')}]: {f.get('step','')}\n"
            f"Output:\n{f.get('stdout','')[:500]}"
        )
    return "\n\n".join(parts)


def report_node(state: AnalysisState) -> AnalysisState:
    llm = get_llm(state["model_key"])  # ← defined here, passed into _ask_llm

    doc = Document()
    successful = [f for f in state["findings"] if f.get("status") == "success"]
    all_findings_text = _format_findings_text(successful)

    # Title
    title_para = doc.add_heading("AI Data Analysis Report", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Question: {state['user_question']}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)

    # Executive Summary
    _add_heading(doc, "Executive Summary", 1)
    exec_text = _ask_llm(
        llm,
        f"Write a 2-3 sentence executive summary answering: '{state['user_question']}'\n"
        f"Based on:\n{all_findings_text[:1500]}\n"
        f"Professional English only. No markdown. No bullet points."
    )
    _add_paragraph(doc, exec_text)

    # Dataset Overview
    _add_heading(doc, "Dataset Overview", 1)
    profile_lines = state["data_profile"].split("\n")
    overview_data = []
    for line in profile_lines[:35]:
        line = line.strip()
        if ":" in line and not line.startswith("-"):
            parts = line.split(":", 1)
            if len(parts) == 2:
                overview_data.append((parts[0].strip(), parts[1].strip()))

    if overview_data:
        tbl = doc.add_table(rows=len(overview_data[:20]), cols=2)
        tbl.style = "Table Grid"
        for i, (key, val) in enumerate(overview_data[:20]):
            tbl.rows[i].cells[0].text = key
            tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            _shade_cell(tbl.rows[i].cells[0], "EBF4FA")
            tbl.rows[i].cells[1].text = val
            tbl.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph("")

    # User Question
    _add_heading(doc, "User Question", 1)
    _add_paragraph(doc, state["user_question"])

    # Data Quality
    _add_heading(doc, "Data Quality", 1)
    quality_findings = [
        f for f in state["findings"]
        if any(kw in f.get("step", "").lower() for kw in ["missing", "duplicate", "quality", "null"])
    ]
    if quality_findings:
        for f in quality_findings:
            _add_heading(doc, f["step"][:80], 2)
            if f["stdout"]:
                _add_table_from_text(doc, f["stdout"])
            if f.get("chart_path"):
                _embed_chart(doc, f["chart_path"])
    else:
        _add_paragraph(doc, "No significant data quality issues were detected.")

    # Key Statistical Findings
    _add_heading(doc, "Key Statistical Findings", 1)
    for f in successful:
        _add_heading(doc, f["step"][:80], 2)
        if f["stdout"]:
            interpretation = _ask_llm(
                llm,
                f"In one professional sentence, summarize what this output means:\n"
                f"{f['stdout'][:500]}\nContext: {f['step']}\nNo markdown."
            )
            _add_paragraph(doc, interpretation)
            _add_table_from_text(doc, f["stdout"])
        if f.get("chart_path"):
            _embed_chart(doc, f["chart_path"])
            cap = doc.add_paragraph(f"Figure: {f['step'][:70]}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True

    # Business Interpretation
    _add_heading(doc, "Business Interpretation", 1)
    biz = _ask_llm(
        llm,
        f"Write 3-4 sentences interpreting the business meaning of these results.\n"
        f"User question: {state['user_question']}\n"
        f"Findings:\n{all_findings_text[:2000]}\n"
        f"Be specific, reference actual numbers. No markdown."
    )
    _add_paragraph(doc, biz)

    # Recommendations
    _add_heading(doc, "Recommendations", 1)
    recs = _ask_llm(
        llm,
        f"Write 4-5 concrete business recommendations based on:\n"
        f"Question: {state['user_question']}\n"
        f"Findings:\n{all_findings_text[:2000]}\n"
        f"Format: numbered list, one sentence each. No markdown headers."
    )
    for line in recs.split("\n"):
        line = re.sub(r"^\d+[\.\)]\s*", "", line.strip())
        if line:
            _add_bullet(doc, line)

    # Conclusion
    _add_heading(doc, "Conclusion", 1)
    conclusion = _ask_llm(
        llm,
        f"Write a 2-3 sentence conclusion.\n"
        f"Analysis answered: {state['user_question']}\n"
        f"Key results: {all_findings_text[:1000]}\n"
        f"Professional tone. No markdown."
    )
    _add_paragraph(doc, conclusion)

    report_path = os.path.join(config.REPORTS_FOLDER, "report.docx")
    doc.save(report_path)
    state["report_path"] = report_path
    print(f"\n=== REPORT SAVED: {report_path} ===\n")
    return state